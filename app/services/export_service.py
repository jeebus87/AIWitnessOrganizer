"""Export service for generating PDF, Excel, and DOCX reports"""
import io
from datetime import datetime
from typing import List, Dict, Any, Optional

import pandas as pd
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph,
    Spacer, PageBreak
)

from app.db.models import Witness, Matter, Document


class ExportService:
    """Service for generating PDF and Excel witness reports"""

    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Set up custom paragraph styles for PDF"""
        self.styles.add(ParagraphStyle(
            name="WitnessName",
            parent=self.styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            textColor=colors.darkblue
        ))

        self.styles.add(ParagraphStyle(
            name="Observation",
            parent=self.styles["Normal"],
            fontSize=8,
            leading=10,
            textColor=colors.black
        ))

        self.styles.add(ParagraphStyle(
            name="CoverTitle",
            parent=self.styles["Title"],
            fontSize=24,
            spaceAfter=20
        ))

    def _format_witness_info(self, w: Dict[str, Any]) -> str:
        """Format witness info combining name, role, and contact"""
        name = w.get("full_name", "Unknown")
        role = w.get("role", "").replace("_", " ").title()

        # Format contact info
        contact_parts = []
        if w.get("address"):
            contact_parts.append(w["address"])
        if w.get("phone"):
            contact_parts.append(w["phone"])
        if w.get("email"):
            contact_parts.append(w["email"])

        if contact_parts:
            contact_str = ", ".join(contact_parts)
        else:
            contact_str = "Contact info unknown at this time"

        return f"{name}, {role}, {contact_str}"

    def _format_source_document(self, w: Dict[str, Any]) -> str:
        """Format source document with page number if available"""
        source_doc = w.get("document_filename", "") or ""
        source_page = w.get("source_page")
        if source_page:
            source_doc = f"{source_doc} (Page {source_page})"
        return source_doc

    def _format_relevance(self, w: Dict[str, Any]) -> str:
        """Format relevance level as user-friendly text with reason"""
        # Map relevance values to display text
        relevance_display = {
            "HIGHLY_RELEVANT": "Highly Relevant",
            "highly_relevant": "Highly Relevant",
            "RELEVANT": "Relevant",
            "relevant": "Relevant",
            "SOMEWHAT_RELEVANT": "Somewhat Relevant",
            "somewhat_relevant": "Somewhat Relevant",
            "NOT_RELEVANT": "Not Relevant",
            "not_relevant": "Not Relevant",
            # Legacy importance values fallback
            "HIGH": "Highly Relevant",
            "high": "Highly Relevant",
            "MEDIUM": "Relevant",
            "medium": "Relevant",
            "LOW": "Somewhat Relevant",
            "low": "Somewhat Relevant",
        }

        # Get relevance (prefer new field, fallback to importance)
        relevance = w.get("relevance") or w.get("importance") or "RELEVANT"
        relevance_text = relevance_display.get(str(relevance), "Relevant")

        # Add reason if available
        relevance_reason = w.get("relevance_reason", "") or ""
        if relevance_reason:
            return f"{relevance_text} - {relevance_reason}"
        return relevance_text

    def _get_relevance_sort_key(self, w: Dict[str, Any]) -> int:
        """Get sort key for relevance (lower = more relevant = first)"""
        relevance = str(w.get("relevance") or w.get("importance") or "RELEVANT").upper().replace(" ", "_")
        relevance_order = {
            "HIGHLY_RELEVANT": 0, "HIGH": 0,
            "RELEVANT": 1, "MEDIUM": 1,
            "SOMEWHAT_RELEVANT": 2, "LOW": 2,
            "NOT_RELEVANT": 3,
        }
        return relevance_order.get(relevance, 1)

    def witnesses_to_dataframe(
        self,
        witnesses: List[Dict[str, Any]],
        include_document_info: bool = True
    ) -> pd.DataFrame:
        """
        Convert witness data to a pandas DataFrame.
        Structure: Witness Info, Relevance, Confidence, Observation, Source Summary, Source Document
        """
        # Sort witnesses by relevance: HIGHLY_RELEVANT first, then RELEVANT, etc.
        sorted_witnesses = sorted(
            witnesses,
            key=lambda w: self._get_relevance_sort_key(w)
        )

        # Group witnesses by name to handle multiple observations
        from collections import OrderedDict
        witness_groups = OrderedDict()
        for w in sorted_witnesses:
            name = w.get("full_name", "Unknown")
            if name not in witness_groups:
                witness_groups[name] = []
            witness_groups[name].append(w)

        rows = []
        for witness_name, observations in witness_groups.items():
            first_obs = observations[0]

            if len(observations) == 1:
                # Single observation - show everything in one row
                w = observations[0]
                row = {
                    "Witness Info": self._format_witness_info(w),
                    "Relevance": self._format_relevance(w),
                    "Confidence": f"{w.get('confidence_score', 0) * 100:.0f}%",
                    "Observation": w.get("observation", "") or "",
                    "Source Summary": w.get("source_quote", "") or "",
                    "Source Document": self._format_source_document(w)
                }
                rows.append(row)
            else:
                # Multiple observations - first row is summary
                summary_observation = f"Multiple observations ({len(observations)} entries) - see details below"

                summary_row = {
                    "Witness Info": self._format_witness_info(first_obs),
                    "Relevance": self._format_relevance(first_obs),
                    "Confidence": f"{first_obs.get('confidence_score', 0) * 100:.0f}%",
                    "Observation": summary_observation,
                    "Source Summary": "See Below",
                    "Source Document": "See Below"
                }
                rows.append(summary_row)

                # Subsequent rows - individual observations
                for w in observations:
                    row = {
                        "Witness Info": "",  # Blank for continuation rows
                        "Relevance": "",
                        "Confidence": "",
                        "Observation": w.get("observation", "") or "",
                        "Source Summary": w.get("source_quote", "") or "",
                        "Source Document": self._format_source_document(w)
                    }
                    rows.append(row)

        return pd.DataFrame(rows)

    def generate_excel(
        self,
        witnesses: List[Dict[str, Any]],
        matter_name: Optional[str] = None,
        matter_number: Optional[str] = None,
        firm_name: Optional[str] = None,
        generated_by: Optional[str] = None,
        include_document_info: bool = True
    ) -> bytes:
        """
        Generate an Excel file with witness data.
        Structure matches PDF: Witness Info, Importance, Confidence, Observation, Source Summary, Source Document

        Returns:
            Excel file as bytes
        """
        df = self.witnesses_to_dataframe(witnesses, include_document_info)

        # Create Excel in memory
        output = io.BytesIO()

        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            # Start data at row 6 to leave room for header info
            header_row_start = 6
            df.to_excel(writer, sheet_name="Witnesses", index=False, startrow=header_row_start)

            workbook = writer.book
            worksheet = writer.sheets["Witnesses"]

            # Formats
            title_format = workbook.add_format({
                "bold": True,
                "font_size": 18,
                "font_color": "#1E3A5F"
            })

            info_format = workbook.add_format({
                "bold": True,
                "font_size": 11
            })

            info_value_format = workbook.add_format({
                "font_size": 11
            })

            header_format = workbook.add_format({
                "bold": True,
                "bg_color": "#1E3A5F",
                "font_color": "white",
                "border": 1,
                "text_wrap": True,
                "valign": "vcenter"
            })

            high_format = workbook.add_format({
                "bg_color": "#FFE6E6",
                "border": 1,
                "text_wrap": True,
                "valign": "top"
            })

            medium_format = workbook.add_format({
                "bg_color": "#FFF9E6",
                "border": 1,
                "text_wrap": True,
                "valign": "top"
            })

            low_format = workbook.add_format({
                "bg_color": "#E6FFE6",
                "border": 1,
                "text_wrap": True,
                "valign": "top"
            })

            continuation_format = workbook.add_format({
                "bg_color": "#F5F5F5",
                "border": 1,
                "text_wrap": True,
                "valign": "top"
            })

            # Write header information
            worksheet.write(0, 0, "Witness Summary Report", title_format)

            row = 1
            if firm_name:
                worksheet.write(row, 0, f"Firm: {firm_name}", info_format)
                row += 1

            if matter_name:
                matter_display = matter_name
                if matter_number and matter_number != matter_name:
                    matter_display = f"{matter_number} - {matter_name}"
                worksheet.write(row, 0, f"Matter: {matter_display}", info_format)
                row += 1

            if generated_by:
                worksheet.write(row, 0, f"Generated by: {generated_by}", info_format)
                row += 1

            worksheet.write(row, 0, f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", info_format)

            # Apply header format to column headers
            for col_num, value in enumerate(df.columns.values):
                worksheet.write(header_row_start, col_num, value, header_format)

            # Track relevance for row coloring
            # We need to track which witness each row belongs to for proper coloring
            if len(df) > 0:
                current_relevance = None
                for row_num in range(len(df)):
                    excel_row = header_row_start + 1 + row_num
                    relevance = df.iloc[row_num]["Relevance"]
                    witness_info = df.iloc[row_num]["Witness Info"]

                    # If this is a new witness row (has witness info), update relevance
                    if witness_info:
                        current_relevance = relevance

                    # Apply formatting based on relevance (check prefix since it may include reason)
                    rel_str = str(current_relevance or "").lower()
                    if rel_str.startswith("highly relevant"):
                        row_format = high_format
                    elif rel_str.startswith("relevant"):
                        row_format = medium_format
                    elif rel_str.startswith("somewhat relevant"):
                        row_format = low_format
                    elif rel_str.startswith("not relevant"):
                        row_format = continuation_format
                    else:
                        row_format = continuation_format

                    # Write each cell with the appropriate format
                    for col_num, col_name in enumerate(df.columns):
                        value = df.iloc[row_num][col_name]
                        worksheet.write(excel_row, col_num, value, row_format)

            # Set column widths - matching PDF proportions
            # Witness Info, Relevance, Confidence, Observation, Source Summary, Source Document
            col_widths = [35, 40, 12, 45, 35, 30]  # Wider Relevance column for reason text
            for idx, width in enumerate(col_widths):
                if idx < len(df.columns):
                    worksheet.set_column(idx, idx, width)

            # Set row height for data rows to accommodate text wrapping
            # Calculate row height based on content length (estimate ~15 chars per line at font size 10)
            if len(df) > 0:
                for row_num in range(len(df)):
                    # Get the max text length in wrappable columns (Observation, Source Summary)
                    obs_text = str(df.iloc[row_num].get("Observation", "") or "")
                    summary_text = str(df.iloc[row_num].get("Source Summary", "") or "")

                    # Estimate lines needed (col width ~40 chars for observation, ~30 for summary)
                    obs_lines = max(1, len(obs_text) // 50 + 1)
                    summary_lines = max(1, len(summary_text) // 40 + 1)
                    max_lines = max(obs_lines, summary_lines)

                    # Set row height: 15 points per line, minimum 30, maximum 200
                    row_height = min(200, max(30, max_lines * 15))
                    worksheet.set_row(header_row_start + 1 + row_num, row_height)

            # Add auto-filter
            if len(df) > 0:
                worksheet.autofilter(header_row_start, 0, header_row_start + len(df), len(df.columns) - 1)

            # Freeze panes - freeze header row
            worksheet.freeze_panes(header_row_start + 1, 0)

        output.seek(0)
        return output.getvalue()

    def generate_pdf(
        self,
        witnesses: List[Dict[str, Any]],
        matter_name: Optional[str] = None,
        matter_number: Optional[str] = None,
        firm_name: Optional[str] = None,
        generated_by: Optional[str] = None,
        include_cover: bool = True
    ) -> bytes:
        """
        Generate a PDF report with witness data.

        Returns:
            PDF file as bytes
        """
        output = io.BytesIO()

        doc = SimpleDocTemplate(
            output,
            pagesize=landscape(LETTER),
            rightMargin=0.5 * inch,
            leftMargin=0.5 * inch,
            topMargin=0.5 * inch,
            bottomMargin=0.5 * inch
        )

        elements = []

        # Cover page
        if include_cover:
            elements.extend(self._create_cover_page(
                matter_name, matter_number, firm_name, generated_by
            ))
            elements.append(PageBreak())

        # Witness table
        elements.extend(self._create_witness_table(witnesses))

        doc.build(elements)
        output.seek(0)
        return output.getvalue()

    def _create_cover_page(
        self,
        matter_name: Optional[str],
        matter_number: Optional[str],
        firm_name: Optional[str] = None,
        generated_by: Optional[str] = None
    ) -> List:
        """Create the cover page elements"""
        elements = []

        # Firm name at top if provided
        if firm_name:
            elements.append(Spacer(1, 0.5 * inch))
            elements.append(Paragraph(
                firm_name,
                self.styles["Heading2"]
            ))
            elements.append(Spacer(1, 1 * inch))
        else:
            elements.append(Spacer(1, 2 * inch))

        # Title
        elements.append(Paragraph(
            "Witness Summary Report",
            self.styles["CoverTitle"]
        ))

        # Matter info
        if matter_name:
            elements.append(Spacer(1, 0.5 * inch))
            elements.append(Paragraph(
                f"<b>Matter:</b> {matter_name}",
                self.styles["Normal"]
            ))

        if matter_number:
            elements.append(Paragraph(
                f"<b>Matter Number:</b> {matter_number}",
                self.styles["Normal"]
            ))

        # Generated by and date
        elements.append(Spacer(1, 0.5 * inch))
        if generated_by:
            elements.append(Paragraph(
                f"<b>Generated by:</b> {generated_by}",
                self.styles["Normal"]
            ))
        elements.append(Paragraph(
            f"<b>Generated:</b> {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
            self.styles["Normal"]
        ))

        # Footer
        elements.append(Spacer(1, 2 * inch))
        elements.append(Paragraph(
            "Generated by AI Witness Finder",
            self.styles["Italic"]
        ))

        return elements

    def _create_witness_table(self, witnesses: List[Dict[str, Any]]) -> List:
        """Create the witness data table"""
        elements = []

        # Header
        elements.append(Paragraph(
            "Identified Witnesses",
            self.styles["Heading1"]
        ))
        elements.append(Spacer(1, 0.25 * inch))

        if not witnesses:
            elements.append(Paragraph(
                "No witnesses were identified in the analyzed documents.",
                self.styles["Normal"]
            ))
            return elements

        # Sort witnesses by relevance: HIGHLY_RELEVANT first, then RELEVANT, etc.
        sorted_witnesses = sorted(
            witnesses,
            key=lambda w: self._get_relevance_sort_key(w)
        )

        # Group witnesses by name to handle multiple observations
        from collections import OrderedDict
        witness_groups = OrderedDict()
        for w in sorted_witnesses:
            name = w.get("full_name", "Unknown")
            if name not in witness_groups:
                witness_groups[name] = []
            witness_groups[name].append(w)

        # Table headers - new structure with Relevance
        headers = ["Witness Info", "Relevance", "Confidence", "Observation", "Source Summary", "Source Document"]

        data = [headers]

        for witness_name, observations in witness_groups.items():
            first_obs = observations[0]

            # Build witness info: Name, Role, Contact
            name = first_obs.get("full_name", "Unknown")
            role = first_obs.get("role", "").replace("_", " ").title()

            # Format contact info
            contact_parts = []
            if first_obs.get("address"):
                contact_parts.append(first_obs["address"])
            if first_obs.get("phone"):
                contact_parts.append(first_obs["phone"])
            if first_obs.get("email"):
                contact_parts.append(first_obs["email"])

            if contact_parts:
                contact_str = ", ".join(contact_parts)
            else:
                contact_str = "Contact info unknown at this time"

            # Combined witness info
            witness_info = f"{name}, {role}, {contact_str}"

            # Handle single vs multiple observations differently
            if len(observations) == 1:
                # Single observation - show everything in one row
                w = observations[0]
                observation = w.get("observation", "") or ""
                source_summary = w.get("source_quote", "") or ""

                source_doc = w.get("document_filename", "") or ""
                source_page = w.get("source_page")
                if source_page:
                    source_doc = f"{source_doc} (Page {source_page})"

                confidence = f"{w.get('confidence_score', 0) * 100:.0f}%"
                relevance_text = self._format_relevance(w)

                row = [
                    Paragraph(witness_info, self.styles["Observation"]),
                    Paragraph(relevance_text, self.styles["Observation"]),
                    confidence,
                    Paragraph(observation, self.styles["Observation"]),
                    Paragraph(source_summary, self.styles["Observation"]),
                    Paragraph(source_doc, self.styles["Observation"])
                ]
                data.append(row)
            else:
                # Multiple observations
                # First row: just indicate multiple observations, details below
                summary_observation = f"Multiple observations ({len(observations)} entries) - see details below"

                first_w = observations[0]
                confidence = f"{first_w.get('confidence_score', 0) * 100:.0f}%"
                relevance_text = self._format_relevance(first_w)

                summary_row = [
                    Paragraph(witness_info, self.styles["Observation"]),
                    Paragraph(relevance_text, self.styles["Observation"]),
                    confidence,
                    Paragraph(summary_observation, self.styles["Observation"]),
                    "See Below",
                    "See Below"
                ]
                data.append(summary_row)

                # Subsequent rows - individual observations
                for w in observations:
                    observation = w.get("observation", "") or ""
                    source_summary = w.get("source_quote", "") or ""

                    source_doc = w.get("document_filename", "") or ""
                    source_page = w.get("source_page")
                    if source_page:
                        source_doc = f"{source_doc} (Page {source_page})"

                    row = [
                        "",  # Witness Info
                        "",  # Relevance
                        "",  # Confidence
                        Paragraph(observation, self.styles["Observation"]),
                        Paragraph(source_summary, self.styles["Observation"]),
                        Paragraph(source_doc, self.styles["Observation"])
                    ]
                    data.append(row)

        # Create table - adjusted column widths for 6 columns
        # Landscape LETTER = 11" wide, minus 1" margins = 10" available
        # Witness Info, Relevance (wider for reason text), Confidence, Observation, Source Summary, Source Document
        col_widths = [1.8 * inch, 1.5 * inch, 0.7 * inch, 2.3 * inch, 1.9 * inch, 1.8 * inch]

        table = Table(data, colWidths=col_widths, repeatRows=1, splitByRow=True)

        # Style the table
        style = TableStyle([
            # Header
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 10),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),

            # Body
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("ALIGN", (1, 1), (1, -1), "LEFT"),  # Relevance left-aligned for readability
            ("ALIGN", (2, 1), (2, -1), "CENTER"),  # Confidence centered (column 2)

            # Grid
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("LINEBELOW", (0, 0), (-1, 0), 2, colors.HexColor("#1E3A5F")),

            # Padding
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ])

        # Add row colors based on relevance - track which rows belong to which witness
        row_num = 1
        for witness_name, observations in witness_groups.items():
            # Get relevance level for coloring
            relevance_key = self._get_relevance_sort_key(observations[0])
            # Calculate number of rows for this witness
            if len(observations) == 1:
                num_rows = 1
            else:
                num_rows = 1 + len(observations)  # Summary row + individual rows

            for _ in range(num_rows):
                if relevance_key == 0:  # HIGHLY_RELEVANT
                    style.add("BACKGROUND", (0, row_num), (-1, row_num), colors.HexColor("#FFE6E6"))
                elif relevance_key == 1:  # RELEVANT
                    style.add("BACKGROUND", (0, row_num), (-1, row_num), colors.HexColor("#FFF9E6"))
                elif relevance_key == 2:  # SOMEWHAT_RELEVANT
                    style.add("BACKGROUND", (0, row_num), (-1, row_num), colors.HexColor("#E6FFE6"))
                else:  # NOT_RELEVANT
                    style.add("BACKGROUND", (0, row_num), (-1, row_num), colors.HexColor("#F5F5F5"))
                row_num += 1

        table.setStyle(style)
        elements.append(table)

        return elements

    def _create_relevancy_section(self, relevancy_data: Dict[str, Any]) -> List:
        """
        Create the relevancy analysis section showing allegations, defenses,
        and which witnesses relate to each claim.
        """
        elements = []

        allegations = relevancy_data.get("allegations", [])
        defenses = relevancy_data.get("defenses", [])
        witness_summary = relevancy_data.get("witness_summary", [])

        if not allegations and not defenses:
            return elements

        # Section header
        elements.append(PageBreak())
        elements.append(Paragraph(
            "Relevancy Analysis",
            self.styles["Heading1"]
        ))
        elements.append(Spacer(1, 0.25 * inch))
        elements.append(Paragraph(
            "This section shows the relationship between identified witnesses and the case allegations/defenses.",
            self.styles["Normal"]
        ))
        elements.append(Spacer(1, 0.25 * inch))

        # Allegations section
        if allegations:
            elements.append(Paragraph("Case Allegations", self.styles["Heading2"]))
            elements.append(Spacer(1, 0.1 * inch))

            allegations_data = [["#", "Allegation", "Linked Witnesses"]]
            for alleg in allegations:
                linked = alleg.get("linked_witnesses", [])
                if linked:
                    witness_text = "\n".join([
                        f"• {w.get('witness_name', 'Unknown')} ({w.get('relationship', 'neutral')})"
                        for w in linked
                    ])
                else:
                    witness_text = "No linked witnesses"

                allegations_data.append([
                    str(alleg.get("number", "")),
                    Paragraph(alleg.get("text", ""), self.styles["Observation"]),
                    Paragraph(witness_text, self.styles["Observation"])
                ])

            alleg_table = Table(
                allegations_data,
                colWidths=[0.5 * inch, 5 * inch, 4 * inch],
                repeatRows=1
            )
            alleg_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#8B0000")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            elements.append(alleg_table)
            elements.append(Spacer(1, 0.25 * inch))

        # Defenses section
        if defenses:
            elements.append(Paragraph("Case Defenses", self.styles["Heading2"]))
            elements.append(Spacer(1, 0.1 * inch))

            defenses_data = [["#", "Defense", "Linked Witnesses"]]
            for defense in defenses:
                linked = defense.get("linked_witnesses", [])
                if linked:
                    witness_text = "\n".join([
                        f"• {w.get('witness_name', 'Unknown')} ({w.get('relationship', 'neutral')})"
                        for w in linked
                    ])
                else:
                    witness_text = "No linked witnesses"

                defenses_data.append([
                    str(defense.get("number", "")),
                    Paragraph(defense.get("text", ""), self.styles["Observation"]),
                    Paragraph(witness_text, self.styles["Observation"])
                ])

            def_table = Table(
                defenses_data,
                colWidths=[0.5 * inch, 5 * inch, 4 * inch],
                repeatRows=1
            )
            def_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#006400")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            elements.append(def_table)
            elements.append(Spacer(1, 0.25 * inch))

        # Witness-Claim Breakdown
        if witness_summary:
            elements.append(Paragraph("Witness Relevancy Breakdown", self.styles["Heading2"]))
            elements.append(Spacer(1, 0.1 * inch))

            breakdown_data = [["Witness", "Relevant To"]]
            for witness in witness_summary:
                claim_links = witness.get("claim_links", [])
                if claim_links:
                    links_text = "\n".join([
                        f"• {link.get('claim_type', '').title()} #{link.get('claim_number', '')} "
                        f"({link.get('relationship', 'neutral')}): {link.get('explanation', '')[:80]}"
                        for link in claim_links
                    ])
                else:
                    links_text = "No specific claim links"

                breakdown_data.append([
                    Paragraph(witness.get("name", "Unknown"), self.styles["WitnessName"]),
                    Paragraph(links_text, self.styles["Observation"])
                ])

            breakdown_table = Table(
                breakdown_data,
                colWidths=[2.5 * inch, 7 * inch],
                repeatRows=1
            )
            breakdown_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            elements.append(breakdown_table)

        return elements

    def generate_pdf_with_relevancy(
        self,
        witnesses: List[Dict[str, Any]],
        relevancy_data: Optional[Dict[str, Any]] = None,
        matter_name: Optional[str] = None,
        matter_number: Optional[str] = None,
        firm_name: Optional[str] = None,
        generated_by: Optional[str] = None,
        include_cover: bool = True
    ) -> bytes:
        """
        Generate a PDF report with witness data and relevancy analysis.

        Args:
            witnesses: List of witness dictionaries
            relevancy_data: Optional dictionary with allegations, defenses, and witness links
            matter_name: Name of the matter
            matter_number: Matter number
            firm_name: Name of the law firm
            generated_by: User who generated the report
            include_cover: Whether to include cover page

        Returns:
            PDF file as bytes
        """
        output = io.BytesIO()

        doc = SimpleDocTemplate(
            output,
            pagesize=landscape(LETTER),
            rightMargin=0.5 * inch,
            leftMargin=0.5 * inch,
            topMargin=0.5 * inch,
            bottomMargin=0.5 * inch
        )

        elements = []

        # Cover page
        if include_cover:
            elements.extend(self._create_cover_page(
                matter_name, matter_number, firm_name, generated_by
            ))
            elements.append(PageBreak())

        # Witness table
        elements.extend(self._create_witness_table(witnesses))

        # Relevancy analysis section (if data provided)
        if relevancy_data:
            elements.extend(self._create_relevancy_section(relevancy_data))

        doc.build(elements)
        output.seek(0)
        return output.getvalue()

    def _set_cell_shading(self, cell, color_hex: str):
        """Set cell background color in DOCX table"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex.lstrip('#'))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def generate_docx(
        self,
        witnesses: List[Dict[str, Any]],
        matter_name: Optional[str] = None,
        matter_number: Optional[str] = None,
        firm_name: Optional[str] = None,
        generated_by: Optional[str] = None,
        include_cover: bool = True
    ) -> bytes:
        """
        Generate a DOCX (Word) report with witness data.
        Same content as PDF, but in editable Word format.

        Args:
            witnesses: List of witness dictionaries
            matter_name: Name of the matter
            matter_number: Matter number
            firm_name: Name of the law firm
            generated_by: User who generated the report
            include_cover: Whether to include cover page

        Returns:
            DOCX file as bytes
        """
        doc = DocxDocument()

        # Set document margins (0.5 inches like PDF)
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.page_width = Inches(11)  # Landscape letter
            section.page_height = Inches(8.5)

        # Cover page
        if include_cover:
            self._add_docx_cover_page(doc, matter_name, matter_number, firm_name, generated_by)
            doc.add_page_break()

        # Witness table
        self._add_docx_witness_table(doc, witnesses)

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    def _add_docx_cover_page(
        self,
        doc: DocxDocument,
        matter_name: Optional[str],
        matter_number: Optional[str],
        firm_name: Optional[str] = None,
        generated_by: Optional[str] = None
    ):
        """Add cover page to DOCX document"""
        # Add spacing at top
        for _ in range(3):
            doc.add_paragraph()

        # Firm name
        if firm_name:
            p = doc.add_paragraph()
            run = p.add_run(firm_name)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(30, 58, 95)  # #1E3A5F
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()

        # Title
        p = doc.add_paragraph()
        run = p.add_run("Witness Summary Report")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(30, 58, 95)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Matter info
        if matter_name:
            p = doc.add_paragraph()
            p.add_run("Matter: ").bold = True
            p.add_run(matter_name)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if matter_number and matter_number != matter_name:
            p = doc.add_paragraph()
            p.add_run("Matter Number: ").bold = True
            p.add_run(matter_number)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Generated by and date
        if generated_by:
            p = doc.add_paragraph()
            p.add_run("Generated by: ").bold = True
            p.add_run(generated_by)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.add_run("Generated: ").bold = True
        p.add_run(datetime.now().strftime('%B %d, %Y at %I:%M %p'))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Footer
        for _ in range(5):
            doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Generated by AI Witness Finder")
        run.italic = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_docx_witness_table(self, doc: DocxDocument, witnesses: List[Dict[str, Any]]):
        """Add witness data table to DOCX document"""
        # Header
        p = doc.add_heading("Identified Witnesses", level=1)
        p.runs[0].font.color.rgb = RGBColor(30, 58, 95)

        if not witnesses:
            doc.add_paragraph("No witnesses were identified in the analyzed documents.")
            return

        # Sort witnesses by relevance
        sorted_witnesses = sorted(
            witnesses,
            key=lambda w: self._get_relevance_sort_key(w)
        )

        # Group witnesses by name
        from collections import OrderedDict
        witness_groups = OrderedDict()
        for w in sorted_witnesses:
            name = w.get("full_name", "Unknown")
            if name not in witness_groups:
                witness_groups[name] = []
            witness_groups[name].append(w)

        # Calculate total rows needed
        total_rows = 1  # Header row
        for witness_name, observations in witness_groups.items():
            if len(observations) == 1:
                total_rows += 1
            else:
                total_rows += 1 + len(observations)  # Summary + individual observations

        # Create table with 6 columns
        table = doc.add_table(rows=total_rows, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        widths = [Inches(1.8), Inches(1.5), Inches(0.7), Inches(2.3), Inches(1.9), Inches(1.8)]
        for idx, width in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = width

        # Add headers
        headers = ["Witness Info", "Relevance", "Confidence", "Observation", "Source Summary", "Source Document"]
        for col_idx, header in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            self._set_cell_shading(cell, "1E3A5F")

        # Relevance colors
        relevance_colors = {
            0: "FFE6E6",  # HIGHLY_RELEVANT - light red
            1: "FFF9E6",  # RELEVANT - light yellow
            2: "E6FFE6",  # SOMEWHAT_RELEVANT - light green
            3: "F5F5F5",  # NOT_RELEVANT - light gray
        }

        # Fill data
        row_idx = 1
        for witness_name, observations in witness_groups.items():
            first_obs = observations[0]
            relevance_key = self._get_relevance_sort_key(first_obs)
            row_color = relevance_colors.get(relevance_key, "FFFFFF")

            if len(observations) == 1:
                # Single observation
                w = observations[0]
                row = table.rows[row_idx]

                row.cells[0].text = self._format_witness_info(w)
                row.cells[1].text = self._format_relevance(w)
                row.cells[2].text = f"{w.get('confidence_score', 0) * 100:.0f}%"
                row.cells[3].text = w.get("observation", "") or ""
                row.cells[4].text = w.get("source_quote", "") or ""
                row.cells[5].text = self._format_source_document(w)

                # Apply styling
                for col_idx in range(6):
                    cell = row.cells[col_idx]
                    self._set_cell_shading(cell, row_color)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)

                row_idx += 1
            else:
                # Multiple observations - summary row first
                row = table.rows[row_idx]

                row.cells[0].text = self._format_witness_info(first_obs)
                row.cells[1].text = self._format_relevance(first_obs)
                row.cells[2].text = f"{first_obs.get('confidence_score', 0) * 100:.0f}%"
                row.cells[3].text = f"Multiple observations ({len(observations)} entries) - see details below"
                row.cells[4].text = "See Below"
                row.cells[5].text = "See Below"

                for col_idx in range(6):
                    cell = row.cells[col_idx]
                    self._set_cell_shading(cell, row_color)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)

                row_idx += 1

                # Individual observation rows
                for w in observations:
                    row = table.rows[row_idx]

                    row.cells[0].text = ""
                    row.cells[1].text = ""
                    row.cells[2].text = ""
                    row.cells[3].text = w.get("observation", "") or ""
                    row.cells[4].text = w.get("source_quote", "") or ""
                    row.cells[5].text = self._format_source_document(w)

                    for col_idx in range(6):
                        cell = row.cells[col_idx]
                        self._set_cell_shading(cell, row_color)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)

                    row_idx += 1

    def generate_docx_with_relevancy(
        self,
        witnesses: List[Dict[str, Any]],
        relevancy_data: Optional[Dict[str, Any]] = None,
        matter_name: Optional[str] = None,
        matter_number: Optional[str] = None,
        firm_name: Optional[str] = None,
        generated_by: Optional[str] = None,
        include_cover: bool = True
    ) -> bytes:
        """
        Generate a DOCX report with witness data and relevancy analysis.

        Args:
            witnesses: List of witness dictionaries
            relevancy_data: Optional dictionary with allegations, defenses, and witness links
            matter_name: Name of the matter
            matter_number: Matter number
            firm_name: Name of the law firm
            generated_by: User who generated the report
            include_cover: Whether to include cover page

        Returns:
            DOCX file as bytes
        """
        doc = DocxDocument()

        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.page_width = Inches(11)  # Landscape letter
            section.page_height = Inches(8.5)

        # Cover page
        if include_cover:
            self._add_docx_cover_page(doc, matter_name, matter_number, firm_name, generated_by)
            doc.add_page_break()

        # Witness table
        self._add_docx_witness_table(doc, witnesses)

        # Relevancy analysis section
        if relevancy_data:
            self._add_docx_relevancy_section(doc, relevancy_data)

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    def _add_docx_relevancy_section(self, doc: DocxDocument, relevancy_data: Dict[str, Any]):
        """Add relevancy analysis section to DOCX document"""
        allegations = relevancy_data.get("allegations", [])
        defenses = relevancy_data.get("defenses", [])
        witness_summary = relevancy_data.get("witness_summary", [])

        if not allegations and not defenses:
            return

        doc.add_page_break()

        # Section header
        p = doc.add_heading("Relevancy Analysis", level=1)
        p.runs[0].font.color.rgb = RGBColor(30, 58, 95)

        doc.add_paragraph(
            "This section shows the relationship between identified witnesses and the case allegations/defenses."
        )

        # Allegations section
        if allegations:
            p = doc.add_heading("Case Allegations", level=2)
            p.runs[0].font.color.rgb = RGBColor(139, 0, 0)  # Dark red

            table = doc.add_table(rows=len(allegations) + 1, cols=3)
            table.style = 'Table Grid'

            # Headers
            headers = ["#", "Allegation", "Linked Witnesses"]
            for col_idx, header in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                self._set_cell_shading(cell, "8B0000")

            # Data rows
            for idx, alleg in enumerate(allegations):
                row = table.rows[idx + 1]
                row.cells[0].text = str(alleg.get("number", ""))

                row.cells[1].text = alleg.get("text", "")

                linked = alleg.get("linked_witnesses", [])
                if linked:
                    witness_text = "\n".join([
                        f"- {w.get('witness_name', 'Unknown')} ({w.get('relationship', 'neutral')})"
                        for w in linked
                    ])
                else:
                    witness_text = "No linked witnesses"
                row.cells[2].text = witness_text

            doc.add_paragraph()

        # Defenses section
        if defenses:
            p = doc.add_heading("Case Defenses", level=2)
            p.runs[0].font.color.rgb = RGBColor(0, 100, 0)  # Dark green

            table = doc.add_table(rows=len(defenses) + 1, cols=3)
            table.style = 'Table Grid'

            # Headers
            headers = ["#", "Defense", "Linked Witnesses"]
            for col_idx, header in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                self._set_cell_shading(cell, "006400")

            # Data rows
            for idx, defense in enumerate(defenses):
                row = table.rows[idx + 1]
                row.cells[0].text = str(defense.get("number", ""))

                row.cells[1].text = defense.get("text", "")

                linked = defense.get("linked_witnesses", [])
                if linked:
                    witness_text = "\n".join([
                        f"- {w.get('witness_name', 'Unknown')} ({w.get('relationship', 'neutral')})"
                        for w in linked
                    ])
                else:
                    witness_text = "No linked witnesses"
                row.cells[2].text = witness_text

            doc.add_paragraph()

        # Witness-Claim Breakdown
        if witness_summary:
            p = doc.add_heading("Witness Relevancy Breakdown", level=2)
            p.runs[0].font.color.rgb = RGBColor(30, 58, 95)

            table = doc.add_table(rows=len(witness_summary) + 1, cols=2)
            table.style = 'Table Grid'

            # Headers
            headers = ["Witness", "Relevant To"]
            for col_idx, header in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                self._set_cell_shading(cell, "1E3A5F")

            # Data rows
            for idx, witness in enumerate(witness_summary):
                row = table.rows[idx + 1]
                row.cells[0].text = witness.get("name", "Unknown")
                row.cells[0].paragraphs[0].runs[0].bold = True
                row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 139)

                claim_links = witness.get("claim_links", [])
                if claim_links:
                    links_text = "\n".join([
                        f"- {link.get('claim_type', '').title()} #{link.get('claim_number', '')} "
                        f"({link.get('relationship', 'neutral')}): {link.get('explanation', '')[:80]}"
                        for link in claim_links
                    ])
                else:
                    links_text = "No specific claim links"
                row.cells[1].text = links_text
